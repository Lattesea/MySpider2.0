#!/usr/bin/python3
# -*- coding: utf-8 -*-


#读取docx中的文本代码示例
import docx
import re

#获取文档
file=docx.Document("王志禄.docx")
print("段落数:"+str(len(file.paragraphs))) #输出段落数
file_word = docx.Document()

#输出每一段的内容
for para in file.paragraphs:
    print(para.text)
print('-------------')
#输出段落编号及段落内容
para_data = []
for i in range(len(file.paragraphs)):
    # for j in map(lambda x:x.split(' '),file.paragraphs[i].text):
    para_single = file.paragraphs[i].text.split(' ')

    while '' in para_single:  # 移除空格
        para_single.remove('')
    if para_single == []:
        pass
    else:
        print("第" + str(i) + "段的内容是：", para_single)
    # para_data.append(para_single)
    # for data_number in range(len(para_single)):
    #     data_num = re.findall(r"\d", para_single[data_number])
    #     data_num = ''.join(data_num)
    #     para_data.append(data_num + '    ')
# file_word.add_paragraph(para_data)
# file_word.save("number.docx")