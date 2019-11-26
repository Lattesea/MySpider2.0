# coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

# 打开文档
document = Document()
# 加入不同等级的标题
document.add_heading(u'MS WORD写入测试', 0)
document.add_heading(u'一级标题', 1)
document.add_heading(u'二级标题', 2)

document.save(u'test.docx')
